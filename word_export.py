import re
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH # Import for paragraph alignment
import base64
import io

def insert_images_to_word_from_markdown(text, image_list, output_path):
    doc = Document()
    
    # Regex mới để bắt cả [HÌNH: img-X.jpeg] và [BẢNG: table-X.jpeg]
    # Nhóm 1: Loại (HÌNH hoặc BẢNG)
    # Nhóm 2: Tên file ảnh (img-X.jpeg hoặc table-X.jpeg)
    pattern = r'\[(HÌNH|BẢNG):\s*(img-\d+\.jpeg|table-\d+\.jpeg)\]'
    
    pos = 0
    
    # Tạo một dictionary để dễ dàng tìm kiếm ảnh theo tên
    image_dict = {img["name"]: img for img in image_list}

    for match in re.finditer(pattern, text):
        start, end = match.span()
        img_type, img_name = match.groups() # img_type sẽ là "HÌNH" hoặc "BẢNG"
        
        # Thêm văn bản trước hình ảnh
        before_img = text[pos:start]
        if before_img.strip():
            doc.add_paragraph(before_img)
        
        # Chèn ảnh/bảng
        img_data = image_dict.get(img_name)
        if img_data:
            try:
                img_bytes = base64.b64decode(img_data["base64"])
                p = doc.add_paragraph()
                run = p.add_run()
                
                # Thêm ảnh vào tài liệu
                # Có thể thêm kiểm tra định dạng ảnh nếu cần (ví dụ: if img_data["name"].endswith(".png"))
                run.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
                
                # Căn giữa hình ảnh
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
                
                # Thêm chú thích (caption)
                caption_text = f"({img_type}: {img_name})"
                caption_p = doc.add_paragraph(caption_text)
                caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
            except Exception as e:
                # Xử lý lỗi nếu không thể chèn ảnh (ví dụ: dữ liệu base64 bị hỏng)
                doc.add_paragraph(f"[Lỗi chèn {img_type} {img_name}: {e}]")
        else:
            doc.add_paragraph(f"[Không tìm thấy {img_type}: {img_name}]")
        
        pos = end
    
    # Thêm phần còn lại của văn bản sau hình ảnh cuối cùng (nếu có)
    if text[pos:].strip():
        doc.add_paragraph(text[pos:])
    
    doc.save(output_path)

