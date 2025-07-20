import streamlit as st
import base64
import io
import os
import re
import tempfile
from PIL import Image, ImageEnhance, ImageFilter, ImageOps, ImageStat
import fitz  # PyMuPDF
import pdf2image
from docx import Document
from docx.shared import Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import requests
from typing import List, Tuple, Optional
import json
import numpy as np
from datetime import datetime

# Try to import OpenCV, fallback if not available
try:
    import cv2
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False

# Page configuration
st.set_page_config(
    page_title="P_OCR PDF AI 2025",
    page_icon="📄",
    layout="wide",
    initial_sidebar_state="expanded"
)

# Custom CSS
st.markdown("""
<style>
    .main-header {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        padding: 2rem;
        border-radius: 10px;
        color: white;
        text-align: center;
        margin-bottom: 2rem;
    }
    .stButton > button {
        background: linear-gradient(90deg, #667eea 0%, #764ba2 100%);
        color: white;
        border: none;
        border-radius: 5px;
        padding: 0.5rem 1rem;
        font-weight: bold;
    }
    .success-box {
        background-color: #d4edda;
        border: 1px solid #c3e6cb;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
    .warning-box {
        background-color: #fff3cd;
        border: 1px solid #ffeaa7;
        border-radius: 5px;
        padding: 1rem;
        margin: 1rem 0;
    }
</style>
""", unsafe_allow_html=True)

class AdvancedImageProcessor:
    """Enhanced image processing with smart content detection"""
    
    @staticmethod
    def analyze_image_content(image: Image.Image) -> dict:
        """Analyze image to determine content characteristics"""
        try:
            # Convert to array for analysis
            img_array = np.array(image.convert('RGB'))
            
            if OPENCV_AVAILABLE:
                gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            else:
                gray = np.array(image.convert('L'))
            
            # Calculate statistics
            analysis = {
                'width': image.width,
                'height': image.height,
                'aspect_ratio': image.width / image.height if image.height > 0 else 1.0,
                'mean_brightness': float(np.mean(gray)),
                'std_brightness': float(np.std(gray)),
                'is_mostly_white': bool(np.mean(gray) > 240),
                'has_high_contrast': bool(np.std(gray) > 50),
                'complexity_score': 0.0,
                'edge_density': 0.0
            }
            
            if OPENCV_AVAILABLE:
                # Edge detection for complexity
                try:
                    edges = cv2.Canny(gray.astype(np.uint8), 50, 150)
                    analysis['edge_density'] = float(np.sum(edges > 0) / (image.width * image.height))
                    analysis['complexity_score'] = analysis['edge_density'] * analysis['std_brightness']
                except Exception:
                    pass
            
            return analysis
            
        except Exception as e:
            st.warning(f"Lỗi phân tích ảnh: {str(e)}")
            return {
                'width': image.width,
                'height': image.height,
                'aspect_ratio': 1.0,
                'mean_brightness': 128.0,
                'std_brightness': 0.0,
                'is_mostly_white': False,
                'has_high_contrast': False,
                'complexity_score': 0.0,
                'edge_density': 0.0
            }
    
    @staticmethod
    def is_illustration_image(image: Image.Image, debug=False) -> bool:
        """Determine if image is likely an illustration/diagram"""
        analysis = AdvancedImageProcessor.analyze_image_content(image)
        
        # Criteria for illustrations
        min_size = 80
        max_aspect_ratio = 8.0
        min_complexity = 0.005
        min_contrast = 15
        
        if debug:
            st.write(f"**Debug ảnh {image.width}x{image.height}:**")
            st.write(f"- Complexity: {analysis['complexity_score']:.4f} (min: {min_complexity})")
            st.write(f"- Contrast: {analysis['std_brightness']:.1f} (min: {min_contrast})")
            st.write(f"- Aspect ratio: {analysis['aspect_ratio']:.2f}")
            st.write(f"- Mean brightness: {analysis['mean_brightness']:.1f}")
        
        # Check criteria
        size_ok = image.width >= min_size and image.height >= min_size
        aspect_ok = 1/max_aspect_ratio <= analysis['aspect_ratio'] <= max_aspect_ratio
        has_content = analysis['complexity_score'] >= min_complexity or analysis['std_brightness'] >= min_contrast
        not_empty = not analysis['is_mostly_white'] or analysis['mean_brightness'] < 250
        
        is_illustration = size_ok and aspect_ok and has_content and not_empty
        
        # Special cases
        if not is_illustration and image.width >= 100 and image.height >= 100 and analysis['std_brightness'] >= 10:
            is_illustration = True
            if debug:
                st.info("✅ Giữ ảnh theo tiêu chí đặc biệt (size + contrast)")
        
        if debug:
            result_text = "✅ GIỮ" if is_illustration else "❌ LOẠI BỎ"
            st.write(f"**Kết quả: {result_text}**")
            st.write("---")
        
        return is_illustration
    
    @staticmethod
    def enhance_image_quality(image: Image.Image) -> Image.Image:
        """Enhanced image quality improvement"""
        try:
            if image.mode not in ['RGB', 'RGBA']:
                image = image.convert('RGB')
            
            analysis = AdvancedImageProcessor.analyze_image_content(image)
            
            # Adjust enhancement parameters
            if analysis['mean_brightness'] < 100:
                brightness_factor = 1.2
                contrast_factor = 1.3
            elif analysis['mean_brightness'] > 200:
                brightness_factor = 0.95
                contrast_factor = 1.1
            else:
                brightness_factor = 1.0
                contrast_factor = 1.2
            
            # Apply enhancements
            if brightness_factor != 1.0:
                enhancer = ImageEnhance.Brightness(image)
                image = enhancer.enhance(brightness_factor)
            
            enhancer = ImageEnhance.Contrast(image)
            image = enhancer.enhance(contrast_factor)
            
            if analysis['std_brightness'] < 40:
                enhancer = ImageEnhance.Sharpness(image)
                image = enhancer.enhance(1.3)
            
            image = ImageOps.autocontrast(image, cutoff=1)
            
            return image
            
        except Exception as e:
            st.warning(f"Không thể cải thiện ảnh: {str(e)}")
            return image
    
    @staticmethod
    def smart_content_crop(image: Image.Image) -> Image.Image:
        """Smart cropping to preserve important content"""
        if not OPENCV_AVAILABLE:
            return AdvancedImageProcessor.simple_content_crop(image)
        
        try:
            img_array = np.array(image.convert('RGB'))
            gray = cv2.cvtColor(img_array, cv2.COLOR_RGB2GRAY)
            
            # Edge detection
            edges1 = cv2.Canny(gray, 50, 150)
            edges2 = cv2.Canny(gray, 30, 100)
            combined_edges = cv2.bitwise_or(edges1, edges2)
            
            # Morphological operations
            kernel = np.ones((3,3), np.uint8)
            edges_dilated = cv2.dilate(combined_edges, kernel, iterations=2)
            edges_closed = cv2.morphologyEx(edges_dilated, cv2.MORPH_CLOSE, kernel)
            
            # Find contours
            contours, _ = cv2.findContours(edges_closed, cv2.RETR_EXTERNAL, cv2.CHAIN_APPROX_SIMPLE)
            
            if not contours:
                return image
            
            # Filter significant contours
            significant_contours = []
            for contour in contours:
                x, y, w, h = cv2.boundingRect(contour)
                area = w * h
                if area > 1000 and w > 50 and h > 50:
                    significant_contours.append((contour, area, x, y, w, h))
            
            if not significant_contours:
                return image
            
            # Calculate bounding box
            if len(significant_contours) == 1:
                _, _, x, y, w, h = significant_contours[0]
            else:
                all_x = [x for _, _, x, y, w, h in significant_contours]
                all_y = [y for _, _, x, y, w, h in significant_contours]
                all_x2 = [x + w for _, _, x, y, w, h in significant_contours]
                all_y2 = [y + h for _, _, x, y, w, h in significant_contours]
                
                x = min(all_x)
                y = min(all_y)
                w = max(all_x2) - x
                h = max(all_y2) - y
            
            # Add padding
            padding_ratio = 0.05
            padding_x = max(10, int(w * padding_ratio))
            padding_y = max(10, int(h * padding_ratio))
            
            x = max(0, x - padding_x)
            y = max(0, y - padding_y)
            w = min(image.width - x, w + 2 * padding_x)
            h = min(image.height - y, h + 2 * padding_y)
            
            # Only crop if significant reduction
            crop_area = w * h
            original_area = image.width * image.height
            
            if crop_area >= 0.5 * original_area:
                return image.crop((x, y, x + w, y + h))
            
            return image
            
        except Exception as e:
            st.warning(f"Smart crop failed: {str(e)}")
            return image
    
    @staticmethod
    def simple_content_crop(image: Image.Image) -> Image.Image:
        """Simple content-based cropping"""
        try:
            gray = image.convert('L')
            img_array = np.array(gray)
            
            content_mask = img_array < 240
            content_coords = np.argwhere(content_mask)
            
            if len(content_coords) == 0:
                return image
            
            y_coords, x_coords = content_coords[:, 0], content_coords[:, 1]
            y1, y2 = y_coords.min(), y_coords.max()
            x1, x2 = x_coords.min(), x_coords.max()
            
            padding = 20
            x1 = max(0, x1 - padding)
            y1 = max(0, y1 - padding)
            x2 = min(image.width, x2 + padding)
            y2 = min(image.height, y2 + padding)
            
            crop_area = (x2 - x1) * (y2 - y1)
            original_area = image.width * image.height
            
            if crop_area >= 0.3 * original_area:
                return image.crop((x1, y1, x2, y2))
            
            return image
            
        except Exception as e:
            st.warning(f"Simple crop failed: {str(e)}")
            return image
    
    @staticmethod
    def process_image_for_ocr(image: Image.Image) -> Image.Image:
        """Complete processing pipeline for OCR"""
        # Enhance quality
        image = AdvancedImageProcessor.enhance_image_quality(image)
        
        # Smart crop if illustration
        if AdvancedImageProcessor.is_illustration_image(image):
            image = AdvancedImageProcessor.smart_content_crop(image)
        
        # Resize if too large
        max_dimension = 1200
        if max(image.width, image.height) > max_dimension:
            if image.width > image.height:
                ratio = max_dimension / image.width
                new_height = int(image.height * ratio)
                image = image.resize((max_dimension, new_height), Image.Resampling.LANCZOS)
            else:
                ratio = max_dimension / image.height
                new_width = int(image.width * ratio)
                image = image.resize((new_width, max_dimension), Image.Resampling.LANCZOS)
        
        return image

class AdvancedFormulaProcessor:
    """Enhanced LaTeX formula detection and processing"""
    
    @staticmethod
    def get_math_patterns() -> List[Tuple[str, str]]:
        """Get comprehensive math patterns"""
        return [
            (r'\b[a-zA-Z]\w*\s*[=]\s*[^=\n]{3,}', 'equation'),
            (r'\b\d+/\d+\b', 'fraction'),
            (r'\b[a-zA-Z]+/[a-zA-Z]+\b', 'variable_fraction'),
            (r'\b[a-zA-Z]\w*\^\d+\b', 'power'),
            (r'\b[a-zA-Z]\w*\^{[^}]+}\b', 'complex_power'),
            (r'√\([^)]+\)', 'sqrt_parentheses'),
            (r'√\d+', 'sqrt_number'),
            (r'√[a-zA-Z]+', 'sqrt_variable'),
            (r'\b(alpha|beta|gamma|delta|epsilon|zeta|eta|theta|iota|kappa|lambda|mu|nu|xi|pi|rho|sigma|tau|upsilon|phi|chi|psi|omega)\b', 'greek'),
            (r'[∫∑∏∆∇±×÷≤≥≠≈∞∈∉⊂⊃∪∩∀∃∄]', 'math_symbols'),
            (r'\b[a-zA-Z]\w*[_\^][a-zA-Z0-9]+\b', 'sub_super'),
            (r'\b[a-zA-Z]\w*[_\^]{[^}]+}\b', 'complex_sub_super'),
            (r'\b(sin|cos|tan|log|ln|exp|sqrt|abs|max|min|lim|int|sum|prod)\s*\([^)]+\)', 'functions'),
            (r'\[[^\[\]]*\]', 'matrix_vector'),
            (r'\b[a-zA-Z]\w*\s*[\+\-\*/]\s*[a-zA-Z0-9\+\-\*/\^\(\)\s]{5,}', 'complex_expression'),
            (r'd[a-zA-Z]/d[a-zA-Z]', 'derivative'),
            (r'∂[a-zA-Z]/∂[a-zA-Z]', 'partial_derivative'),
        ]
    
    @staticmethod
    def clean_formula_content(content: str) -> str:
        """Clean and normalize formula content"""
        content = re.sub(r'\s+', ' ', content.strip())
        
        replacements = {
            'х': 'x', 'у': 'y', '—': '-', '–': '-', '×': '*', '÷': '/'
        }
        
        for old, new in replacements.items():
            content = content.replace(old, new)
        
        return content
    
    @staticmethod
    def wrap_math_formulas(text: str) -> str:
        """Enhanced formula wrapping"""
        patterns = AdvancedFormulaProcessor.get_math_patterns()
        processed_text = text
        wrapped_regions = []
        
        for pattern, pattern_type in patterns:
            matches = list(re.finditer(pattern, processed_text, re.IGNORECASE))
            
            for match in reversed(matches):
                start_pos = match.start()
                end_pos = match.end()
                formula = match.group()
                
                # Check if already wrapped
                already_wrapped = any(
                    start_pos >= region[0] and end_pos <= region[1] 
                    for region in wrapped_regions
                )
                
                if already_wrapped:
                    continue
                
                before = processed_text[:start_pos]
                after = processed_text[end_pos:]
                
                # Skip if already in LaTeX
                if ('${' in before[-10:] and '}$' in after[:10]) or \
                   ('$' in before[-5:] and '$' in after[:5]):
                    continue
                
                # Validate math expressions
                if pattern_type in ['equation', 'complex_expression']:
                    if not re.search(r'[=\+\-\*/\^]', formula):
                        continue
                    
                    if before.endswith(' ') and after.startswith(' ') and \
                       not re.search(r'[.!?]\s*$', before):
                        sentence_before = before.split('.')[-1] if '.' in before else before
                        math_keywords = ['equation', 'formula', 'calculate', 'solve', 'result', 'answer']
                        if not any(keyword in sentence_before.lower() for keyword in math_keywords):
                            continue
                
                # Wrap formula
                clean_formula = AdvancedFormulaProcessor.clean_formula_content(formula)
                wrapped_formula = f"${{{clean_formula}}}$"
                processed_text = (processed_text[:start_pos] + 
                                wrapped_formula + 
                                processed_text[end_pos:])
                
                # Update regions
                len_diff = len(wrapped_formula) - len(formula)
                wrapped_regions.append((start_pos, end_pos + len_diff))
                wrapped_regions = [(s + len_diff if s > start_pos else s, 
                                  e + len_diff if e > start_pos else e) 
                                 for s, e in wrapped_regions]
        
        return processed_text

class OCRProcessor:
    def __init__(self):
        self.mistral_api_key = None
        self.image_processor = AdvancedImageProcessor()
        self.formula_processor = AdvancedFormulaProcessor()
        
    def setup_api(self, mistral_key: str):
        """Setup Mistral API key"""
        self.mistral_api_key = mistral_key
    
    def extract_images_from_pdf(self, pdf_file, enhance=True, debug=False) -> List[Image.Image]:
        """Extract high-quality illustrations from PDF"""
        images = []
        rejected_images = []
        
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(pdf_file.read())
                tmp_path = tmp_file.name
            
            pdf_document = fitz.open(tmp_path)
            
            if debug:
                st.write("**🔍 Debug: Trích xuất ảnh từ PDF**")
            
            for page_num in range(pdf_document.page_count):
                page = pdf_document[page_num]
                image_list = page.get_images()
                
                if debug:
                    st.write(f"**Trang {page_num + 1}: Tìm thấy {len(image_list)} ảnh**")
                
                for img_index, img in enumerate(image_list):
                    try:
                        xref = img[0]
                        pix = fitz.Pixmap(pdf_document, xref)
                        
                        if pix.n - pix.alpha < 4:  # GRAY or RGB
                            img_data = pix.tobytes("png")
                            img_pil = Image.open(io.BytesIO(img_data))
                            
                            if debug:
                                st.write(f"**Ảnh {img_index + 1} (trang {page_num + 1}):**")
                            
                            # Filter images
                            if self.image_processor.is_illustration_image(img_pil, debug=debug):
                                if enhance:
                                    original_size = (img_pil.width, img_pil.height)
                                    img_pil = self.image_processor.process_image_for_ocr(img_pil)
                                    new_size = (img_pil.width, img_pil.height)
                                    if debug:
                                        st.success(f"✅ Đã xử lý: {original_size} → {new_size}")
                                
                                images.append(img_pil)
                                if debug:
                                    st.success(f"✅ ĐÃ GIỮ ảnh {len(images)}")
                            else:
                                rejected_images.append((img_pil, page_num + 1, img_index + 1))
                                if debug:
                                    st.error(f"❌ ĐÃ LOẠI BỎ")
                        
                        pix = None
                    except Exception as e:
                        if debug:
                            st.warning(f"Lỗi xử lý ảnh {img_index} trang {page_num + 1}: {str(e)}")
                        continue
            
            pdf_document.close()
            os.unlink(tmp_path)
            
            # Summary
            if debug:
                st.write("**📊 Tổng kết trích xuất:**")
                st.success(f"✅ **Giữ lại: {len(images)} ảnh**")
                if rejected_images:
                    st.error(f"❌ **Loại bỏ: {len(rejected_images)} ảnh**")
            
        except Exception as e:
            st.error(f"Lỗi trích xuất ảnh từ PDF: {str(e)}")
        
        return images
    
    def convert_pdf_to_images(self, pdf_file) -> List[Image.Image]:
        """Convert PDF pages to images for OCR"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.pdf') as tmp_file:
                tmp_file.write(pdf_file.read())
                tmp_path = tmp_file.name
            
            pages = pdf2image.convert_from_path(tmp_path, dpi=300)
            os.unlink(tmp_path)
            
            return pages
        except Exception as e:
            st.error(f"Lỗi chuyển đổi PDF: {str(e)}")
            return []
    
    def ocr_with_mistral(self, image: Image.Image) -> str:
        """OCR with Mistral AI"""
        if not self.mistral_api_key:
            return "Mistral API key không được cung cấp"
        
        try:
            buffer = io.BytesIO()
            image.save(buffer, format='PNG')
            img_b64 = base64.b64encode(buffer.getvalue()).decode()
            
            headers = {
                'Authorization': f'Bearer {self.mistral_api_key}',
                'Content-Type': 'application/json'
            }
            
            prompt = """Trích xuất văn bản từ hình ảnh với độ chính xác tối đa. YÊU CẦU:

1. **Văn bản thường**: Nhận diện chính xác 100% tất cả chữ

2. **Công thức toán học**: 
   - Bọc MỌI công thức bằng ${...}$ 
   - VD: ${x^2 + y^2 = z^2}$, ${a/b}$, ${√x}$, ${∫f(x)dx}$

3. **Hình ảnh/biểu đồ/sơ đồ**: 
   - KHI GẶP bất kỳ ảnh/biểu đồ/sơ đồ nào, PHẢI ghi:
   ![Hình 1](image1.png)
   ![Hình 2](image2.png)
   - Đánh số tuần tự 1, 2, 3...
   - Đặt marker ở đúng vị trí ảnh xuất hiện

4. **Định dạng**: Giữ nguyên xuống dòng, thụt lề

QUAN TRỌNG:
- MỌI công thức toán PHẢI có ${...}$
- MỌI ảnh minh họa PHẢI có ![Hình X](imageX.png)
- Không bỏ sót nội dung nào"""
            
            payload = {
                "model": "mistral-small-latest",
                "temperature": 0.1,
                "max_tokens": 4000,
                "messages": [
                    {
                        "role": "user",
                        "content": [
                            {"type": "text", "text": prompt},
                            {"type": "image_url", "image_url": {"url": f"data:image/png;base64,{img_b64}"}}
                        ]
                    }
                ]
            }
            
            response = requests.post(
                "https://api.mistral.ai/v1/chat/completions",
                headers=headers,
                json=payload,
                timeout=60
            )
            
            if response.status_code == 200:
                result = response.json()
                extracted_text = result['choices'][0]['message']['content']
                processed_text = self.formula_processor.wrap_math_formulas(extracted_text)
                return processed_text
            else:
                return f"Lỗi Mistral API: {response.status_code} - {response.text}"
            
        except Exception as e:
            return f"Lỗi Mistral API: {str(e)}"

class WordExporter:
    def __init__(self):
        self.doc = Document()
        self.setup_document_style()
    
    def setup_document_style(self):
        """Setup document styling"""
        section = self.doc.sections[0]
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
    
    def add_content(self, text: str, images: List[Image.Image] = None):
        """Add content with precise image placement"""
        # Title
        title = self.doc.add_heading('Kết quả OCR - P_OCR PDF AI 2025', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Timestamp
        timestamp = self.doc.add_paragraph(f'Ngày tạo: {datetime.now().strftime("%d/%m/%Y %H:%M:%S")}')
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Main content
        self.doc.add_heading('Nội dung văn bản được trích xuất:', level=1)
        self._process_content_with_images(text, images or [])
    
    def _process_content_with_images(self, text: str, images: List[Image.Image]):
        """Process content with precise image placement"""
        # Find all image markers
        image_marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
        all_markers = re.findall(image_marker_pattern, text, re.IGNORECASE)
        
        st.write(f"**📊 Word Export Debug:**")
        st.write(f"- Số ảnh có sẵn: {len(images)}")
        st.write(f"- Image markers tìm thấy: {all_markers}")
        
        # Process line by line
        lines = text.split('\n')
        
        for line_num, line in enumerate(lines):
            if not line.strip():
                self.doc.add_paragraph()
                continue
            
            # Check for image markers in this line
            matches = list(re.finditer(image_marker_pattern, line, re.IGNORECASE))
            
            if matches:
                st.write(f"🎯 Line {line_num}: Tìm thấy {len(matches)} markers")
                
                current_pos = 0
                for match in matches:
                    # Add text before marker
                    before_text = line[current_pos:match.start()].strip()
                    if before_text:
                        p = self.doc.add_paragraph()
                        self._add_text_with_math_formatting(p, before_text)
                    
                    # Process image marker
                    image_num = int(match.group(1))
                    marker_text = match.group(0)
                    
                    if 1 <= image_num <= len(images):
                        img = images[image_num - 1]
                        success = self._insert_image_safely(img, f"Hình {image_num}")
                        if success:
                            st.success(f"✅ Đã chèn Hình {image_num}")
                        else:
                            st.error(f"❌ Lỗi chèn Hình {image_num}")
                    else:
                        # Image not available
                        p = self.doc.add_paragraph()
                        run = p.add_run(f"[🖼️ Hình {image_num} - Không có ảnh tương ứng]")
                        run.italic = True
                        run.bold = True
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                    current_pos = match.end()
                
                # Add remaining text
                remaining_text = line[current_pos:].strip()
                if remaining_text:
                    p = self.doc.add_paragraph()
                    self._add_text_with_math_formatting(p, remaining_text)
            else:
                # Regular line without images
                p = self.doc.add_paragraph()
                self._add_text_with_math_formatting(p, line.strip())
        
        # Add unreferenced images
        self._add_unreferenced_images(text, images)
    
    def _add_text_with_math_formatting(self, paragraph, text: str):
        """Add text with LaTeX formatting"""
        if not text:
            return
            
        # Split by LaTeX formulas
        latex_pattern = r'(\$\{[^}]+\}\$)'
        parts = re.split(latex_pattern, text)
        
        for part in parts:
            if part.startswith('${') and part.endswith('}$'):
                # Math formula
                formula_content = part[2:-2]  # Remove ${ and }$
                run = paragraph.add_run(formula_content)
                run.bold = True
                run.italic = True
            else:
                # Regular text
                if part:
                    paragraph.add_run(part)
    
    def _insert_image_safely(self, img: Image.Image, caption: str) -> bool:
        """Insert image with error handling"""
        try:
            with tempfile.NamedTemporaryFile(delete=False, suffix='.png') as tmp_img:
                # Convert to RGB if needed
                if img.mode not in ['RGB']:
                    img = img.convert('RGB')
                
                # Save image
                img.save(tmp_img.name, 'PNG', quality=95)
                
                # Calculate size
                max_width_cm = 12.0
                img_width_cm = (img.width / 96) * 2.54
                img_height_cm = (img.height / 96) * 2.54
                
                scale = min(max_width_cm / img_width_cm, 1.0) if img_width_cm > max_width_cm else 1.0
                final_width = Cm(img_width_cm * scale)
                
                # Create paragraph and insert image
                p = self.doc.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add breaks
                p.add_run().add_break()
                
                # Insert image
                run = p.add_run()
                run.add_picture(tmp_img.name, width=final_width)
                
                # Add caption
                p.add_run().add_break()
                caption_run = p.add_run(f"({caption})")
                caption_run.italic = True
                caption_run.bold = True
                
                p.add_run().add_break()
                
                # Clean up
                os.unlink(tmp_img.name)
                
                return True
                
        except Exception as e:
            st.error(f"Lỗi chèn ảnh {caption}: {str(e)}")
            # Add placeholder
            try:
                p = self.doc.add_paragraph()
                run = p.add_run(f"[{caption}: Lỗi chèn ảnh]")
                run.italic = True
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except:
                pass
            return False
    
    def _add_unreferenced_images(self, text: str, images: List[Image.Image]):
        """Add images not referenced in text"""
        # Find referenced image numbers
        referenced_nums = set()
        pattern = r'!\[Hình\s*(\d+)\]'
        for match in re.finditer(pattern, text):
            referenced_nums.add(int(match.group(1)))
        
        # Find unreferenced images
        unreferenced = []
        for i, img in enumerate(images, 1):
            if i not in referenced_nums:
                unreferenced.append((i, img))
        
        if unreferenced:
            self.doc.add_page_break()
            self.doc.add_heading('Hình ảnh bổ sung (chưa được tham chiếu):', level=1)
            
            for img_num, img in unreferenced:
                self._insert_image_safely(img, f"Hình {img_num}")
    
    def add_statistics(self, images: List[Image.Image], text: str):
        """Add statistics section"""
        self.doc.add_page_break()
        self.doc.add_heading('Thống kê chi tiết:', level=1)
        
        # Calculate stats
        word_count = len(text.split())
        char_count = len(text)
        
        # Count formulas and image references
        formula_count = len(re.findall(r'\$\{[^}]+\}\
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        # Image stats
        if images:
            avg_width = sum(img.width for img in images) / len(images)
            avg_height = sum(img.height for img in images) / len(images)
            total_pixels = sum(img.width * img.height for img in images)
        else:
            avg_width = avg_height = total_pixels = 0
        
        # Create table
        stats_table = self.doc.add_table(rows=8, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ['Tổng số từ:', f'{word_count:,}'],
            ['Tổng ký tự:', f'{char_count:,}'],
            ['Số đoạn văn:', f'{paragraphs}'],
            ['Công thức LaTeX:', f'{formula_count}'],
            ['Hình ảnh trích xuất:', f'{len(images)}'],
            ['Tham chiếu ảnh:', f'{image_refs}'],
            ['Kích thước ảnh TB:', f'{avg_width:.0f}×{avg_height:.0f}px' if images else 'N/A'],
            ['Tổng pixel ảnh:', f'{total_pixels:,}' if images else 'N/A']
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
    
    def save(self) -> bytes:
        """Save document and return bytes"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# Global patterns for validation
FORMULA_PATTERN = r'(\$\{[^}]+\}\$)'  # Added capturing group
IMAGE_MARKER_PATTERN = r'(!\[Hình\s*\d+\])'  # Added capturing group

def count_formulas(text: str) -> int:
    """Count LaTeX formulas in text"""
    return len(re.findall(r'\$\{[^}]+\}\

def main():
    """Main application"""
    # Prevent multiple initializations
    if 'main_running' in st.session_state and st.session_state.main_running:
        return
    
    st.session_state.main_running = True
    
    try:
        # Header
        st.markdown("""
        <div class="main-header">
            <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
            <p>Ứng dụng OCR thông minh với Mistral AI</p>
            <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
        </div>
        """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            # Use capturing groups for replacement
            highlighted_text = re.sub(r'(\$\{[^}]+\}\$)', r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Use non-capturing pattern for finding
                formula_matches = re.findall(r'\$\{[^}]+\}\
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    # Ensure main only runs once per session
    if 'app_initialized' not in st.session_state:
        st.session_state.app_initialized = True
    main()
, text))  # Use non-capturing pattern for counting

def count_image_markers(text: str) -> int:
    """Count image markers in text"""
    return len(re.findall(r'!\[Hình\s*\d+\]', text))  # Use non-capturing pattern for counting

def main():
    """Main application"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
        <p>Ứng dụng OCR thông minh với Mistral AI</p>
        <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            highlighted_text = re.sub(FORMULA_PATTERN, r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                formula_matches = re.findall(FORMULA_PATTERN, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text))  # Use non-capturing pattern for counting

def count_image_markers(text: str) -> int:
    """Count image markers in text"""
    return len(re.findall(r'!\[Hình\s*\d+\]', text))  # Use non-capturing pattern for counting

def main():
    """Main application"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
        <p>Ứng dụng OCR thông minh với Mistral AI</p>
        <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            highlighted_text = re.sub(FORMULA_PATTERN, r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                formula_matches = re.findall(FORMULA_PATTERN, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text))
        image_refs = len(re.findall(r'!\[Hình\s*\d+\]', text))
        paragraphs = len([p for p in text.split('\n') if p.strip()])
        
        # Image stats
        if images:
            avg_width = sum(img.width for img in images) / len(images)
            avg_height = sum(img.height for img in images) / len(images)
            total_pixels = sum(img.width * img.height for img in images)
        else:
            avg_width = avg_height = total_pixels = 0
        
        # Create table
        stats_table = self.doc.add_table(rows=8, cols=2)
        stats_table.style = 'Table Grid'
        
        stats_data = [
            ['Tổng số từ:', f'{word_count:,}'],
            ['Tổng ký tự:', f'{char_count:,}'],
            ['Số đoạn văn:', f'{paragraphs}'],
            ['Công thức LaTeX:', f'{formula_count}'],
            ['Hình ảnh trích xuất:', f'{len(images)}'],
            ['Tham chiếu ảnh:', f'{image_refs}'],
            ['Kích thước ảnh TB:', f'{avg_width:.0f}×{avg_height:.0f}px' if images else 'N/A'],
            ['Tổng pixel ảnh:', f'{total_pixels:,}' if images else 'N/A']
        ]
        
        for i, (label, value) in enumerate(stats_data):
            stats_table.cell(i, 0).text = label
            stats_table.cell(i, 1).text = value
    
    def save(self) -> bytes:
        """Save document and return bytes"""
        buffer = io.BytesIO()
        self.doc.save(buffer)
        buffer.seek(0)
        return buffer.getvalue()

# Global patterns for validation
FORMULA_PATTERN = r'(\$\{[^}]+\}\$)'  # Added capturing group
IMAGE_MARKER_PATTERN = r'(!\[Hình\s*\d+\])'  # Added capturing group

def count_formulas(text: str) -> int:
    """Count LaTeX formulas in text"""
    return len(re.findall(r'\$\{[^}]+\}\

def main():
    """Main application"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
        <p>Ứng dụng OCR thông minh với Mistral AI</p>
        <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            # Use capturing groups for replacement
            highlighted_text = re.sub(r'(\$\{[^}]+\}\$)', r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                # Use non-capturing pattern for finding
                formula_matches = re.findall(r'\$\{[^}]+\}\
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text))  # Use non-capturing pattern for counting

def count_image_markers(text: str) -> int:
    """Count image markers in text"""
    return len(re.findall(r'!\[Hình\s*\d+\]', text))  # Use non-capturing pattern for counting

def main():
    """Main application"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
        <p>Ứng dụng OCR thông minh với Mistral AI</p>
        <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            highlighted_text = re.sub(FORMULA_PATTERN, r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                formula_matches = re.findall(FORMULA_PATTERN, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
, text))  # Use non-capturing pattern for counting

def count_image_markers(text: str) -> int:
    """Count image markers in text"""
    return len(re.findall(r'!\[Hình\s*\d+\]', text))  # Use non-capturing pattern for counting

def main():
    """Main application"""
    # Header
    st.markdown("""
    <div class="main-header">
        <h1>📄 P_OCR PDF AI 2025 - Enhanced Edition</h1>
        <p>Ứng dụng OCR thông minh với Mistral AI</p>
        <p>🎯 Cắt ảnh thông minh • 💯 LaTeX chính xác • 📍 Chèn đúng vị trí</p>
    </div>
    """, unsafe_allow_html=True)
    
    # Initialize processor
    if 'ocr_processor' not in st.session_state:
        st.session_state.ocr_processor = OCRProcessor()
    
    # Sidebar
    with st.sidebar:
        st.header("🔧 Cấu hình")
        
        mistral_key = st.text_input(
            "Mistral API Key",
            type="password",
            help="Nhập API key của Mistral AI"
        )
        
        if st.button("💾 Lưu cấu hình"):
            st.session_state.ocr_processor.setup_api(mistral_key)
            st.success("✅ Đã lưu cấu hình!")
        
        st.markdown("---")
        
        # Options
        st.header("🖼️ Tùy chọn xử lý")
        enhance_images = st.checkbox(
            "Cải thiện chất lượng ảnh",
            value=True,
            help="Tự động cắt thông minh và cải thiện ảnh"
        )
        
        debug_mode = st.checkbox(
            "Chế độ debug",
            value=False,
            help="Hiển thị thông tin debug chi tiết"
        )
        
        st.markdown("---")
        
        # Features info
        st.header("✨ Tính năng nâng cao")
        st.markdown("""
        **🎯 Cắt ảnh thông minh:**
        - Phân biệt ảnh minh họa vs trang trí
        - Bảo toàn nội dung quan trọng
        - Loại bỏ border và noise
        
        **💯 LaTeX chính xác:**
        - Nhận diện công thức phức tạp
        - Xử lý ký hiệu toán học
        - Tránh false positives
        
        **📍 Chèn ảnh đúng vị trí:**
        - Marker `![Hình X](imageX.png)`
        - Resize thông minh
        - Layout chuyên nghiệp
        """)
        
        # Session info
        if 'extracted_images' in st.session_state:
            st.markdown("---")
            st.subheader("📊 Session hiện tại")
            st.write(f"**Ảnh minh họa**: {len(st.session_state.extracted_images)}")
            if 'extracted_text' in st.session_state:
                formula_count = count_formulas(st.session_state.extracted_text)
                image_refs = count_image_markers(st.session_state.extracted_text)
                st.write(f"**Công thức LaTeX**: {formula_count}")
                st.write(f"**Tham chiếu ảnh**: {image_refs}")
    
    # Main content
    col1, col2 = st.columns([2, 1])
    
    with col1:
        st.header("📁 Upload File")
        
        uploaded_file = st.file_uploader(
            "Chọn file PDF hoặc hình ảnh",
            type=['pdf', 'png', 'jpg', 'jpeg'],
            help="Hỗ trợ file PDF và hình ảnh"
        )
        
        if uploaded_file is not None:
            file_type = uploaded_file.type
            st.info(f"📄 File: {uploaded_file.name} ({file_type})")
            
            if st.button("🚀 Bắt đầu OCR nâng cao", type="primary"):
                if not mistral_key:
                    st.error("❌ Vui lòng nhập Mistral API Key!")
                    return
                
                with st.spinner("🔄 Đang xử lý với AI thông minh..."):
                    try:
                        extracted_text = ""
                        extracted_images = []
                        
                        if file_type == "application/pdf":
                            st.info("📄 Đang phân tích PDF...")
                            
                            # Extract illustrations
                            pdf_file_copy = io.BytesIO(uploaded_file.read())
                            extracted_images = st.session_state.ocr_processor.extract_images_from_pdf(
                                pdf_file_copy, enhance=enhance_images, debug=debug_mode
                            )
                            
                            if len(extracted_images) == 0:
                                st.warning("⚠️ Không tìm thấy ảnh minh họa nào!")
                            else:
                                st.success(f"🖼️ Đã trích xuất {len(extracted_images)} ảnh minh họa")
                            
                            # Convert pages for OCR
                            uploaded_file.seek(0)
                            page_images = st.session_state.ocr_processor.convert_pdf_to_images(uploaded_file)
                            
                            # OCR each page
                            st.info(f"🔍 Bắt đầu OCR {len(page_images)} trang...")
                            progress_bar = st.progress(0)
                            
                            for i, page_img in enumerate(page_images):
                                st.info(f"🔍 OCR trang {i+1}/{len(page_images)}...")
                                
                                if enhance_images:
                                    page_img = st.session_state.ocr_processor.image_processor.enhance_image_quality(page_img)
                                
                                page_text = st.session_state.ocr_processor.ocr_with_mistral(page_img)
                                
                                # Check for markers
                                page_markers = count_image_markers(page_text)
                                if page_markers > 0:
                                    st.success(f"✅ Trang {i+1}: Tìm thấy {page_markers} image markers")
                                
                                extracted_text += f"\n--- Trang {i+1} ---\n{page_text}\n"
                                progress_bar.progress((i + 1) / len(page_images))
                        
                        else:
                            st.info("🖼️ Đang xử lý hình ảnh với AI...")
                            image = Image.open(uploaded_file)
                            
                            if enhance_images:
                                image = st.session_state.ocr_processor.image_processor.process_image_for_ocr(image)
                            
                            extracted_text = st.session_state.ocr_processor.ocr_with_mistral(image)
                        
                        # Store results
                        st.session_state.extracted_text = extracted_text
                        st.session_state.extracted_images = extracted_images
                        
                        st.success("✅ Hoàn thành OCR nâng cao!")
                        
                        # Analysis
                        formula_count = count_formulas(extracted_text)
                        image_refs = count_image_markers(extracted_text)
                        
                        # Check mismatches
                        if len(extracted_images) > 0 and image_refs == 0:
                            st.warning("⚠️ **Có ảnh nhưng không có image markers trong text!**")
                        elif len(extracted_images) != image_refs and image_refs > 0:
                            st.warning(f"⚠️ **Số lượng không khớp**: {len(extracted_images)} ảnh vs {image_refs} markers")
                        
                        # Success metrics
                        if formula_count > 0:
                            st.success(f"🔢 Đã nhận diện {formula_count} công thức LaTeX!")
                        
                        if image_refs > 0:
                            st.success(f"📍 Đã đánh dấu {image_refs} vị trí chèn ảnh!")
                        
                        if len(extracted_images) > 0 and image_refs > 0 and len(extracted_images) == image_refs:
                            st.success("🎯 **Perfect match**: Số ảnh và markers khớp hoàn toàn!")
                        
                    except Exception as e:
                        st.error(f"❌ Lỗi: {str(e)}")
    
    with col2:
        st.header("📊 Thống kê nâng cao")
        
        if 'extracted_text' in st.session_state:
            text = st.session_state.extracted_text
            
            word_count = len(text.split())
            char_count = len(text)
            formula_count = count_formulas(text)
            image_refs = count_image_markers(text)
            
            col_a, col_b = st.columns(2)
            with col_a:
                st.metric("Số từ", word_count)
                st.metric("Công thức LaTeX", formula_count)
            with col_b:
                st.metric("Số ký tự", char_count)
                st.metric("Tham chiếu ảnh", image_refs)
            
            if 'extracted_images' in st.session_state:
                st.metric("Ảnh minh họa", len(st.session_state.extracted_images))
                
                # Quality indicators
                if formula_count > 0:
                    st.success("💯 LaTeX detected")
                if image_refs > 0:
                    st.success("📍 Images referenced")
    
    # Results section
    if 'extracted_text' in st.session_state:
        st.markdown("---")
        st.header("📋 Kết quả OCR nâng cao")
        
        # Text display
        with st.expander("📝 Văn bản với LaTeX và markers", expanded=True):
            text = st.session_state.extracted_text
            
            # Highlight text
            highlighted_text = text
            highlighted_text = re.sub(FORMULA_PATTERN, r'**\1**', highlighted_text)
            highlighted_text = re.sub(r'(!\[Hình\s*\d+\]\([^)]*\))', r'***\1***', highlighted_text)
            
            st.text_area(
                "Nội dung (LaTeX: **bold**, Image markers: ***bold-italic***):",
                highlighted_text,
                height=400,
                disabled=True
            )
            
            # Analysis
            col1, col2, col3 = st.columns(3)
            
            with col1:
                formula_matches = re.findall(FORMULA_PATTERN, text)
                st.write("**Công thức LaTeX:**")
                for i, formula in enumerate(formula_matches[:3], 1):
                    st.write(f"{i}. `{formula}`")
                if len(formula_matches) > 3:
                    st.write(f"... và {len(formula_matches) - 3} công thức khác")
            
            with col2:
                image_matches = re.findall(r'!\[Hình\s*(\d+)\]', text)
                st.write("**Markers ảnh:**")
                for img_num in image_matches[:3]:
                    st.write(f"• Hình {img_num}")
                if len(image_matches) > 3:
                    st.write(f"... và {len(image_matches) - 3} ảnh khác")
            
            with col3:
                st.write("**Chất lượng OCR:**")
                st.write(f"✅ {len(text.split())} từ")
                st.write(f"🔢 {len(formula_matches)} công thức")
                st.write(f"📷 {len(image_matches)} refs")
                
                # Validation
                if 'extracted_images' in st.session_state:
                    img_count = len(st.session_state.extracted_images)
                    marker_count = len(image_matches)
                    if img_count == marker_count and img_count > 0:
                        st.success("🎯 Perfect!")
                    elif marker_count == 0:
                        st.error("❌ No markers!")
                    elif img_count != marker_count:
                        st.warning(f"⚠️ Mismatch!")
        
        # Images display
        if 'extracted_images' in st.session_state and st.session_state.extracted_images:
            with st.expander(f"🖼️ Ảnh minh họa ({len(st.session_state.extracted_images)} ảnh)", expanded=False):
                for i, img in enumerate(st.session_state.extracted_images):
                    col1, col2 = st.columns([1, 2])
                    
                    with col1:
                        st.image(img, caption=f"Hình {i+1}", use_column_width=True)
                    
                    with col2:
                        analysis = st.session_state.ocr_processor.image_processor.analyze_image_content(img)
                        st.write(f"**Hình {i+1}:**")
                        st.write(f"• Size: {img.width}×{img.height}px")
                        st.write(f"• Ratio: {analysis['aspect_ratio']:.2f}")
                        st.write(f"• Complexity: {analysis['complexity_score']:.3f}")
                        st.write(f"• Contrast: {analysis['std_brightness']:.1f}")
                    
                    st.markdown("---")
        
        # Export section
        st.markdown("---")
        st.header("📤 Xuất Word chuyên nghiệp")
        
        # Pre-export analysis
        if 'extracted_text' in st.session_state and 'extracted_images' in st.session_state:
            text = st.session_state.extracted_text
            images = st.session_state.extracted_images
            
            marker_pattern = r'!\[Hình\s*(\d+)\]\([^)]*\)'
            found_markers = re.findall(marker_pattern, text, re.IGNORECASE)
            
            col1, col2 = st.columns(2)
            with col1:
                st.write("**📊 Thống kê xuất:**")
                st.write(f"• Ảnh có sẵn: {len(images)}")
                st.write(f"• Markers tìm thấy: {len(found_markers)}")
            
            with col2:
                st.write("**🎯 Tình trạng:**")
                if len(images) == len(found_markers) and len(found_markers) > 0:
                    st.success("✅ Perfect match!")
                elif len(found_markers) == 0:
                    st.warning("⚠️ Không có markers")
                elif len(images) < len(found_markers):
                    st.error("❌ Thiếu ảnh")
                else:
                    st.info("ℹ️ Thừa ảnh")
        
        col1, col2 = st.columns(2)
        with col1:
            include_stats = st.checkbox("Thêm thống kê chi tiết", value=True)
        with col2:
            st.info("🚀 Xuất với tính năng nâng cao")
        
        if st.button("📄 Tạo Word nâng cao", type="primary"):
            with st.spinner("📝 Đang tạo file Word..."):
                try:
                    exporter = WordExporter()
                    
                    images_to_export = st.session_state.get('extracted_images', [])
                    exporter.add_content(st.session_state.extracted_text, images_to_export)
                    
                    if include_stats:
                        exporter.add_statistics(images_to_export, st.session_state.extracted_text)
                    
                    word_bytes = exporter.save()
                    
                    st.success("🎉 File Word đã được tạo thành công!")
                    
                    # Metrics
                    col_a, col_b, col_c, col_d = st.columns(4)
                    with col_a:
                        st.metric("Văn bản", f"{len(st.session_state.extracted_text.split())} từ")
                    with col_b:
                        st.metric("LaTeX", count_formulas(st.session_state.extracted_text))
                    with col_c:
                        st.metric("Ảnh chèn", len(images_to_export))
                    with col_d:
                        st.metric("File size", f"{len(word_bytes)/1024:.1f} KB")
                    
                    # Download
                    filename = f"OCR_Enhanced_{uploaded_file.name.split('.')[0]}.docx"
                    st.download_button(
                        label="⬇️ Tải file Word (Enhanced)",
                        data=word_bytes,
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                    
                    st.success("✅ **Tính năng đã áp dụng**: Cắt ảnh thông minh • LaTeX chính xác • Chèn đúng vị trí")
                    
                except Exception as e:
                    st.error(f"❌ Lỗi tạo Word: {str(e)}")
    
    # Footer
    st.markdown("---")
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 2rem;">
        <p>🚀 <strong>P_OCR PDF AI 2025 - Enhanced Edition</strong></p>
        <p>🎯 <strong>Smart Crop</strong> • 💯 <strong>LaTeX Accuracy</strong> • 📍 <strong>Precise Placement</strong></p>
        <p>🤖 Powered by Mistral AI & Advanced Image Processing</p>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()
