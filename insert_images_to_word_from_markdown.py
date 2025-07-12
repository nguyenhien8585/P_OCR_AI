from docx import Document
from docx.shared import Inches
import base64
import io
import re

def insert_images_to_word_from_markdown(text, image_list, output_path):
    doc = Document()
    # Nếu text có marker ![...](img-x.jpeg) thì chèn đúng vị trí
    pattern = r'!\[([^\]]*)\]\(([^)]+)\)'
    matches = list(re.finditer(pattern, text))
    if matches:
        pos = 0
        for match in matches:
            start, end = match.span()
            caption, img_name = match.groups()
            doc.add_paragraph(text[pos:start])
            found = False
            for img in image_list:
                if img["name"] == img_name:
                    img_bytes = base64.b64decode(img["base64"])
                    doc.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
                    if caption:
                        doc.add_paragraph(f"(Hình: {caption})")
                    found = True
                    break
            if not found:
                doc.add_paragraph(f"[Không tìm thấy ảnh: {img_name}]")
            pos = end
        doc.add_paragraph(text[pos:])
    else:
        # Nếu là ảnh (không có marker), chèn toàn bộ ảnh vào cuối text
        doc.add_paragraph(text)
        if image_list:
            doc.add_paragraph("\nMinh hoạ (trích từ ảnh gốc):")
            for idx, img in enumerate(image_list):
                img_bytes = base64.b64decode(img["base64"])
                doc.add_picture(io.BytesIO(img_bytes), width=Inches(3.5))
                doc.add_paragraph(f"(Hình minh hoạ {idx+1}: {img['name']})")
    doc.save(output_path)
