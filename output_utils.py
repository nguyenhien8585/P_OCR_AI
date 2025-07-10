from docx import Document
from docx.shared import Inches
import tempfile
from io import BytesIO
from PIL import Image

def export_to_word(image_content_pairs):
    doc = Document()
    for img, content in image_content_pairs:
        doc.add_picture(image_to_stream(img), width=Inches(5))
        doc.add_paragraph(content)
        doc.add_page_break()
    output_path = tempfile.NamedTemporaryFile(delete=False, suffix=".docx").name
    doc.save(output_path)
    return output_path

def export_to_latex(image_content_pairs):
    lines = [
        "\\documentclass{article}",
        "\\usepackage{graphicx}",
        "\\begin{document}"
    ]
    for idx, (img, content) in enumerate(image_content_pairs):
        img_path = f"figure{idx+1}.png"
        img.save(img_path)
        lines.append(f"\\begin{{figure}}[h]")
        lines.append(f"\\centering")
        lines.append(f"\\includegraphics[width=0.8\\textwidth]{{{img_path}}}")
        lines.append(f"\\end{{figure}}")
        lines.append(content)
        lines.append("\\newpage")

    lines.append("\\end{document}")
    tex_path = tempfile.NamedTemporaryFile(delete=False, suffix=".tex").name
    with open(tex_path, "w", encoding="utf-8") as f:
        f.write("\n".join(lines))
    return tex_path

def image_to_stream(image: Image.Image):
    buf = BytesIO()
    image.save(buf, format='PNG')
    buf.seek(0)
    return buf
