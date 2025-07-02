def build_docx(text, output_path):
    from docx import Document
    from docx.shared import Inches
    import os
    doc = Document()
    for line in text.splitlines():
        if line.strip().startswith("<<image_"):
            img_name = line.strip().replace("<<", "").replace(">>", "")
            img_path = os.path.join("diagrams", img_name)
            if os.path.exists(img_path):
                doc.add_picture(img_path, width=Inches(4))
        elif line.strip():
            doc.add_paragraph(line)
    doc.save(output_path)

def build_latex(content, output_path):
    preamble = r"""
\documentclass{article}
\usepackage{amsmath,graphicx}
\begin{document}
"""
    with open(output_path, "w", encoding="utf-8") as f:
        f.write(preamble)
        f.write(content)
        f.write("\n\\end{document}")
